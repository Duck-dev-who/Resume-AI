from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import os

class ResumeParserError(Exception):
    """Custom exception for resume parsing errors."""
    pass

def parse_docx_resume(file_path: str) -> str:
    print("Parser received:", repr(file_path))
    print("Ends with .docx?", file_path.lower().endswith(".docx"))
    """
    Parse a DOCX resume and return extracted text.
    Includes paragraphs and tables.
    
    Args:
        file_path (str): Path to the .docx file.
    
    Returns:
        str: Extracted text.
    
    Raises:
        ResumeParserError: If file is invalid or cannot be read.
    """
    if not os.path.exists(file_path):
        raise ResumeParserError(f"File not found: {file_path}")

    if not file_path.lower().endswith(".docx"):
        raise ResumeParserError("Only .docx files are supported.")

    try:
        doc = Document(file_path)
        full_text = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # Extract tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

    except PackageNotFoundError:
        raise ResumeParserError("Invalid or corrupted DOCX file.")
    except Exception as e:
        raise ResumeParserError(str(e))